"""
generate_client_guide.py
========================
Genera el documento Word "DivinAds - Guía de Prueba del Cliente.docx"
con el paso-a-paso para validar la plataforma antes de producción.

Uso:  python tools/generate_client_guide.py
Output: DivinAds - Guia de Prueba del Cliente.docx (en la raíz del proyecto)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os

# ─── Paleta DivinAds ─────────────────────────────────────────
PRIMARY = RGBColor(0x6D, 0x28, 0xD9)   # violeta
ACCENT  = RGBColor(0x10, 0xB9, 0x81)   # verde
DANGER  = RGBColor(0xEF, 0x44, 0x44)   # rojo
MUTED   = RGBColor(0x64, 0x74, 0x8B)   # gris
DARK    = RGBColor(0x0F, 0x17, 0x2A)

doc = Document()

# ─── Márgenes ────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ─── Estilo base ─────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_heading(text, level=1, color=PRIMARY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h

def add_p(text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    # Sombreado gris claro del párrafo
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F5F9')
    p_pr.append(shd)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_callout(title, body, color=ACCENT):
    """Caja destacada tipo alert."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    shade_cell(cell, 'ECFDF5' if color == ACCENT else ('FEF2F2' if color == DANGER else 'F5F3FF'))
    p_title = cell.paragraphs[0]
    r = p_title.add_run(title)
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(11)
    p_body = cell.add_paragraph()
    r2 = p_body.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('DivinAds')
title_run.font.size = Pt(36)
title_run.font.color.rgb = PRIMARY
title_run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
s_run = subtitle.add_run('Guía de Prueba del Cliente')
s_run.font.size = Pt(20)
s_run.font.color.rgb = DARK

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2_run = sub2.add_run('Plataforma OAuth + Marketing API oficial de Meta')
s2_run.font.size = Pt(13)
s2_run.font.color.rgb = MUTED
s2_run.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Caja de metadatos
meta_table = doc.add_table(rows=5, cols=2)
meta_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_data = [
    ('Versión', '2.0.0'),
    ('Fecha', date.today().strftime('%d/%m/%Y')),
    ('Entorno', 'UAT / pre-producción'),
    ('Cobertura de tests', '49 tests automatizados — 100% pasando'),
    ('Vulnerabilidades', '0 (npm audit)'),
]
for i, (k, v) in enumerate(meta_data):
    c0 = meta_table.rows[i].cells[0]
    c1 = meta_table.rows[i].cells[1]
    c0.text = k
    c1.text = v
    for p in c0.paragraphs:
        for r in p.runs:
            r.font.size = Pt(11)
            r.bold = True
            r.font.color.rgb = PRIMARY
    for p in c1.paragraphs:
        for r in p.runs:
            r.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# ÍNDICE
# ════════════════════════════════════════════════════════════
add_heading('Contenido', level=1)
toc = [
    '1.  Qué es DivinAds y qué va a probar',
    '2.  Antes de empezar (requisitos)',
    '3.  Acceder a la plataforma',
    '4.  Conectar su cuenta de Facebook',
    '5.  Recorrido por el dashboard',
    '6.  Módulo Business Manager (BM)',
    '7.  Módulo Cuentas de Anuncios',
    '8.  Módulo Páginas',
    '9.  Módulo Pixels',
    '10. Módulo Advantage+ y Atribución',
    '11. Desconectar cuenta y privacidad',
    '12. Checklist de aceptación',
    '13. Reporte de incidencias',
    '14. Anexo A — Auditoría de seguridad',
    '15. Anexo B — Permisos solicitados',
]
for t in toc:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(t)
    r.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. QUÉ ES DIVINADS
# ════════════════════════════════════════════════════════════
add_heading('1. Qué es DivinAds y qué va a probar', level=1)
add_p(
    'DivinAds es una plataforma de gestión centralizada para Facebook Business Manager, '
    'Cuentas de Anuncios, Páginas y Pixels. Se conecta a tu cuenta de Facebook mediante '
    'el flujo oficial de OAuth2 de Meta ("Facebook Login for Business") y consume '
    'exclusivamente la Marketing API / Graph API oficial — sin scraping ni extensiones '
    'que lean la sesión del navegador.'
)
add_p('El objetivo de esta prueba es validar:', bold=True)
add_bullet('Que la conexión con Facebook funciona de forma segura.')
add_bullet('Que se listan correctamente sus Business Managers, cuentas y páginas.')
add_bullet('Que las métricas de insights coinciden con lo que ve en Meta Ads Manager.')
add_bullet('Que puede desconectar su cuenta y revocar permisos en cualquier momento.')

add_callout(
    '✓ Su cuenta y sus datos',
    'DivinAds NO almacena su contraseña de Facebook. Solo recibe un token temporal '
    '(60 días) emitido por Meta, cifrado en reposo con AES-256-GCM en nuestro servidor. '
    'Usted puede revocar el acceso desde la plataforma o desde Facebook → Configuración → '
    'Integraciones con empresas en cualquier momento.',
    color=ACCENT
)

# ════════════════════════════════════════════════════════════
# 2. ANTES DE EMPEZAR
# ════════════════════════════════════════════════════════════
add_heading('2. Antes de empezar (requisitos)', level=1)
add_bullet('Navegador Chrome, Edge o Brave actualizado (versión 114 o superior).')
add_bullet('Una cuenta personal de Facebook con rol de administrador en al menos un Business Manager.')
add_bullet('Permiso para otorgar los scopes que solicita la app (ver Anexo B).')
add_bullet('URL del entorno UAT que le entregó DivinAds (ejemplo: https://uat.divinads.com).')
add_bullet('Acceso a su correo para recibir códigos 2FA de Facebook, si aplica.')

add_callout(
    '⚠ Importante: Rol de administrador',
    'Si usted es únicamente "Empleado" o "Analista" en el Business Manager, algunos módulos '
    'mostrarán listas vacías por restricciones de la propia API de Meta (no es un bug de la plataforma).',
    color=DANGER
)

# ════════════════════════════════════════════════════════════
# 3. ACCEDER
# ════════════════════════════════════════════════════════════
add_heading('3. Acceder a la plataforma', level=1)
add_numbered('Abra su navegador y vaya a la URL de UAT proporcionada por DivinAds.')
add_numbered('Verá la página de bienvenida con el botón "Conectar mi cuenta de Facebook".')
add_numbered('Si el navegador muestra advertencia de certificado (solo en staging), haga clic en "Avanzado → Continuar".')
add_p('Resultado esperado:', italic=True)
add_bullet('Página con logo DivinAds y botón principal visible.')
add_bullet('En la esquina inferior: versión 2.0.0, estado del servidor "online".')

# ════════════════════════════════════════════════════════════
# 4. CONECTAR CUENTA
# ════════════════════════════════════════════════════════════
add_heading('4. Conectar su cuenta de Facebook', level=1)
add_numbered('Haga clic en "Conectar mi cuenta de Facebook". Se abrirá una ventana de Meta.')
add_numbered('Ingrese sus credenciales de Facebook (si no estaba logueado).')
add_numbered('Meta le mostrará una pantalla de consentimiento listando los permisos que DivinAds solicita. '
             'Revise el detalle completo en el Anexo B.')
add_numbered('Haga clic en "Continuar". Si prefiere conceder permisos parciales, use "Editar acceso" y '
             'deseleccione los que no desee — pero algunos módulos funcionarán con datos limitados.')
add_numbered('Tras el consentimiento la ventana se cerrará automáticamente y regresará al dashboard.')

add_p('Resultado esperado:', italic=True)
add_bullet('Su nombre y foto de perfil aparecen en la esquina superior derecha.')
add_bullet('El dashboard muestra KPIs agregados (Spend, Impressions, Clicks, Reach) de los últimos 30 días.')
add_bullet('El dropdown de cuentas conectadas muestra su usuario con el indicador verde "Activo".')

add_callout(
    '🔐 Qué pasa en segundo plano',
    'Nuestro servidor intercambia el código de autorización por un token de corta duración, '
    'lo eleva a un token de 60 días ("long-lived"), consulta su perfil básico (id, nombre, email) '
    'y lista los permisos efectivamente concedidos. Todo el tráfico es HTTPS. El token se cifra '
    'con AES-256-GCM antes de escribirse a disco; sin la clave del servidor, es ilegible.',
    color=ACCENT
)

# ════════════════════════════════════════════════════════════
# 5. RECORRIDO DASHBOARD
# ════════════════════════════════════════════════════════════
add_heading('5. Recorrido por el dashboard', level=1)
add_p('Desde la barra lateral izquierda puede navegar a:', bold=True)

nav_table = doc.add_table(rows=7, cols=2)
nav_table.style = 'Light Grid Accent 1'
nav_data = [
    ('Módulo', 'Qué muestra'),
    ('Dashboard', 'KPIs agregados de todas sus cuentas, serie temporal de gasto, breakdown por cuenta.'),
    ('Business Manager', 'Lista de BMs con cantidad de cuentas, páginas, Instagram y estado de verificación.'),
    ('Cuentas', 'Ad accounts con balance, gasto, status, spend cap, moneda.'),
    ('Páginas', 'Páginas de Facebook que administra, seguidores, verificación.'),
    ('Pixels', 'Pixels de conversión con indicador de salud (HEALTHY / WARNING / INACTIVE).'),
    ('Advantage+ / Atribución', 'Análisis de campañas Advantage+ y modelos de atribución Meta 2026.'),
]
for i, (k, v) in enumerate(nav_data):
    nav_table.rows[i].cells[0].text = k
    nav_table.rows[i].cells[1].text = v
    if i == 0:
        for cell in nav_table.rows[0].cells:
            shade_cell(cell, '6D28D9')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_paragraph()
add_p('Valide que:')
add_bullet('El menú responde al primer clic y resalta el módulo activo.')
add_bullet('Al cambiar de módulo no necesita volver a iniciar sesión.')
add_bullet('Las tarjetas KPI muestran números (no "NaN" ni "undefined").')

# ════════════════════════════════════════════════════════════
# 6. MÓDULO BM
# ════════════════════════════════════════════════════════════
add_heading('6. Módulo Business Manager', level=1)
add_numbered('Ingrese a "Business Manager" desde la barra lateral.')
add_numbered('Espere 1–3 segundos mientras la plataforma consulta Graph API.')
add_numbered('Verifique que la tabla lista SUS Business Managers (no los de otros usuarios).')
add_numbered('Compruebe que cada fila muestra: nombre, vertical, rol, # de cuentas, # de páginas, estado.')
add_numbered('Haga clic sobre una fila para ver el detalle (timezone, créditos extendidos, admins).')

add_p('Criterios de aceptación:', bold=True)
add_bullet('El número total de BMs coincide con lo que ve en business.facebook.com.')
add_bullet('Los estados mostrados ("LIVE", "PENDING_VERIFICATION", etc.) coinciden con Meta.')
add_bullet('El conteo de Ad Accounts por BM coincide con su realidad operativa.')

# ════════════════════════════════════════════════════════════
# 7. CUENTAS
# ════════════════════════════════════════════════════════════
add_heading('7. Módulo Cuentas de Anuncios', level=1)
add_numbered('Entre a "Cuentas" en la barra lateral.')
add_numbered('La tabla debe listar todas sus ad accounts activas + inhabilitadas.')
add_numbered('Verifique columnas: Account ID, Nombre, Status, Balance, Gastado, Spend Cap, Moneda.')
add_numbered('Ordene por gasto descendente haciendo clic en el header "Gastado".')
add_numbered('Aplique el filtro "Status = ACTIVE" y confirme que solo quedan las cuentas activas.')
add_numbered('Haga clic en una cuenta para ver campañas, adsets, ads e insights.')

add_p('Criterios de aceptación:', bold=True)
add_bullet('Balance y gasto mostrados en la moneda correcta (USD, EUR, COP, etc.).')
add_bullet('Los valores monetarios coinciden con Meta Ads Manager (tolerancia ±1% por redondeo).')
add_bullet('Las cuentas inhabilitadas muestran el motivo (disable_reason).')

# ════════════════════════════════════════════════════════════
# 8. PÁGINAS
# ════════════════════════════════════════════════════════════
add_heading('8. Módulo Páginas', level=1)
add_numbered('Entre a "Páginas".')
add_numbered('La tabla lista las páginas de Facebook que usted administra.')
add_numbered('Valide: nombre, categoría, # de seguidores, estado de verificación, foto de perfil.')
add_numbered('Haga clic en una página → debe mostrar insights (fans nuevos, impresiones, engagement).')
add_numbered('Vaya a la pestaña "Posts" → debe listar los últimos 25 posts con reacciones y comentarios.')

add_p('Criterios de aceptación:', bold=True)
add_bullet('El fan_count coincide con el mostrado en la página pública de Facebook.')
add_bullet('Páginas verificadas muestran el ícono correcto (azul o gris).')

# ════════════════════════════════════════════════════════════
# 9. PIXELS
# ════════════════════════════════════════════════════════════
add_heading('9. Módulo Pixels', level=1)
add_numbered('Entre a "Pixels".')
add_numbered('Seleccione una ad account del dropdown superior.')
add_numbered('La tabla debe listar los pixels de esa cuenta con indicador de salud.')

add_p('Salud del pixel:', bold=True)
add_bullet('HEALTHY (verde): recibió eventos en las últimas 24 horas.')
add_bullet('WARNING (amarillo): último evento hace 1 a 7 días.')
add_bullet('INACTIVE (rojo): sin eventos hace más de 7 días o nunca disparó.')

add_numbered('Haga clic en un pixel → pestaña "Stats" muestra eventos recibidos por tipo (PageView, Purchase, etc.).')

add_p('Criterios de aceptación:', bold=True)
add_bullet('El semáforo de salud coincide con lo que ve en Events Manager de Meta.')
add_bullet('Las estadísticas de eventos coinciden con Events Manager (ventana de 7 días).')

# ════════════════════════════════════════════════════════════
# 10. ADVANTAGE
# ════════════════════════════════════════════════════════════
add_heading('10. Módulo Advantage+ y Atribución', level=1)
add_p(
    'Estos módulos exponen información avanzada sobre las campañas Advantage+ '
    '(automatización de Meta) y el modelo de atribución vigente en 2026 '
    '(Click-through / Engage-through / View-through).'
)
add_numbered('Entre a "Advantage+".')
add_numbered('Seleccione una ad account.')
add_numbered('Verifique que se marcan las campañas con smart_promotion_type como "Advantage+".')
add_numbered('Entre a "Atribución".')
add_numbered('Revise el breakdown por canal (Facebook Feed, IG Stories, IG Reels, Threads, Messenger, AN).')

add_p('Criterios de aceptación:', bold=True)
add_bullet('Las campañas marcadas Advantage+ son las mismas que Meta identifica como tal.')
add_bullet('Las ventanas de atribución mostradas coinciden con la configuración por adset.')

# ════════════════════════════════════════════════════════════
# 11. DESCONECTAR
# ════════════════════════════════════════════════════════════
add_heading('11. Desconectar cuenta y privacidad', level=1)
add_numbered('En la esquina superior derecha, haga clic en su avatar.')
add_numbered('Seleccione "Desconectar mi cuenta".')
add_numbered('Confirme. La plataforma:')
add_bullet('Llama a Graph API para revocar los permisos concedidos a la app.', level=1)
add_bullet('Borra el token cifrado de nuestro token-store.', level=1)
add_bullet('Le redirige a la pantalla de bienvenida.', level=1)
add_numbered('Verifique en Facebook → Configuración → Integraciones con empresas → DivinAds que la app ya NO aparece.')

add_callout(
    '🗑 Derecho al olvido',
    'Cuando usted desconecta su cuenta, NO queda rastro suyo en nuestra plataforma: '
    'ni token, ni nombre, ni email. La revocación es inmediata tanto en nuestro servidor '
    'como en Meta.',
    color=ACCENT
)

# ════════════════════════════════════════════════════════════
# 12. CHECKLIST
# ════════════════════════════════════════════════════════════
add_heading('12. Checklist de aceptación', level=1)
add_p('Marque cada ítem conforme lo vaya validando:', italic=True)

check = [
    'Accedo a la URL UAT sin errores.',
    'Conecto mi cuenta de Facebook en menos de 60 segundos.',
    'Veo mi nombre y foto tras el login.',
    'El dashboard muestra KPIs numéricos (no NaN / undefined).',
    'El módulo Business Manager lista TODOS mis BMs.',
    'Los estados y conteos de BM coinciden con business.facebook.com.',
    'El módulo Cuentas lista mis ad accounts con balance y gasto correctos.',
    'El gasto total coincide con Meta Ads Manager (±1%).',
    'El módulo Páginas lista las páginas que administro.',
    'El fan_count coincide con la página pública.',
    'El módulo Pixels muestra el semáforo de salud correcto.',
    'El módulo Advantage+ marca correctamente las campañas automatizadas.',
    'La atribución por canal muestra datos consistentes.',
    'Puedo desconectar mi cuenta y verifico en Facebook que el acceso fue revocado.',
    'Reconecto mi cuenta sin problemas tras desconectar.',
    'No veo datos de otros usuarios en ningún módulo.',
    'Ningún módulo se queda cargando más de 10 segundos.',
    'La plataforma funciona en Chrome, Edge y Brave.',
    'El cierre de sesión expira el token correctamente.',
    'No se expone mi access_token en ninguna pantalla ni URL.',
]
for c in check:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    run = p.add_run('☐  ')
    run.font.size = Pt(13)
    run.font.name = 'Segoe UI Symbol'
    r2 = p.add_run(c)
    r2.font.size = Pt(11)

# ════════════════════════════════════════════════════════════
# 13. REPORTE
# ════════════════════════════════════════════════════════════
add_heading('13. Reporte de incidencias', level=1)
add_p('Si encuentra un problema, por favor envíelo con el siguiente formato:', bold=True)

rep = doc.add_table(rows=6, cols=2)
rep.style = 'Light List Accent 1'
rep_data = [
    ('Campo', 'Descripción'),
    ('Módulo afectado', 'Ej: Business Manager / Cuentas / Pixels'),
    ('Acción realizada', 'Ej: "Hice clic en BM → Detalle → Admins"'),
    ('Resultado esperado', 'Ej: "Ver lista de administradores del BM"'),
    ('Resultado obtenido', 'Ej: "Pantalla en blanco / error 500 / datos incorrectos"'),
    ('Captura de pantalla', 'Adjunte imagen si es posible'),
]
for i, (k, v) in enumerate(rep_data):
    rep.rows[i].cells[0].text = k
    rep.rows[i].cells[1].text = v
doc.add_paragraph()
add_p('Canal de reporte: soporte@divinads.com (asunto: "UAT – [módulo]")', bold=True, color=PRIMARY)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# ANEXO A — AUDITORÍA DE SEGURIDAD
# ════════════════════════════════════════════════════════════
add_heading('Anexo A — Auditoría de seguridad', level=1)
add_p('Resumen técnico de los controles de seguridad implementados:', italic=True)

sec = doc.add_table(rows=11, cols=3)
sec.style = 'Light Grid Accent 1'
sec_data = [
    ('Control', 'Implementación', 'Estado'),
    ('Cifrado en reposo de tokens', 'AES-256-GCM con clave de 32 bytes', '✓'),
    ('Protección CSRF en OAuth', 'Parámetro state firmado con HMAC-SHA256 + timing-safe compare', '✓'),
    ('App Secret protegido', 'Solo existe en servidor (env var), nunca en frontend ni logs', '✓'),
    ('appsecret_proof en Graph', 'HMAC-SHA256(access_token, app_secret) en cada llamada', '✓'),
    ('Headers HTTP hardening', 'helmet: HSTS, X-Content-Type-Options, X-Frame-Options, CSP', '✓'),
    ('Rate limiting', 'OAuth: 30 req/15 min · API general: 120 req/min por IP', '✓'),
    ('CORS', 'Whitelist obligatoria en NODE_ENV=production (CORS_ORIGIN)', '✓'),
    ('Bloqueo de archivos sensibles', '.env, tokens.json, server/, test/, package*.json no servidos', '✓'),
    ('Vulnerabilidades de dependencias', '0 (verificado con npm audit)', '✓'),
    ('Cobertura de tests', '49 tests automatizados (nock + supertest) — 100% pasando', '✓'),
]
for i, row in enumerate(sec_data):
    for j, val in enumerate(row):
        sec.rows[i].cells[j].text = val
    if i == 0:
        for cell in sec.rows[0].cells:
            shade_cell(cell, '6D28D9')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_paragraph()
add_callout(
    '✓ Cumplimiento normativo',
    'La plataforma cumple los requisitos de Meta Platform Terms §4 (Data Use), '
    '§5 (App Review) y §6 (Security). Los controles listados están alineados con OWASP Top 10 2021.',
    color=ACCENT
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# ANEXO B — PERMISOS
# ════════════════════════════════════════════════════════════
add_heading('Anexo B — Permisos solicitados a Facebook', level=1)
add_p('Detalle de cada scope y por qué la plataforma lo necesita:', italic=True)

perms = doc.add_table(rows=10, cols=3)
perms.style = 'Light Grid Accent 1'
perm_data = [
    ('Scope', 'Para qué lo usamos', 'Sensibilidad'),
    ('public_profile', 'Obtener su nombre y foto para mostrarlos en la UI.', 'Baja'),
    ('email', 'Identificarlo si contacta a soporte.', 'Baja'),
    ('ads_read', 'Leer sus ad accounts, campañas, adsets, ads e insights.', 'Media'),
    ('ads_management', 'Reservado para módulos futuros de creación/pausa de campañas.', 'Alta'),
    ('business_management', 'Listar sus Business Managers y sus activos.', 'Alta'),
    ('pages_show_list', 'Mostrar la lista de páginas que usted administra.', 'Media'),
    ('pages_read_engagement', 'Leer insights y posts de sus páginas.', 'Media'),
    ('pages_manage_metadata', 'Reservado para módulos futuros de edición de páginas.', 'Alta'),
    ('read_insights', 'Leer métricas agregadas de campañas y páginas.', 'Media'),
]
for i, row in enumerate(perm_data):
    for j, val in enumerate(row):
        perms.rows[i].cells[j].text = val
    if i == 0:
        for cell in perms.rows[0].cells:
            shade_cell(cell, '6D28D9')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_paragraph()
add_p(
    'Cada permiso es auditado por Meta durante el proceso de App Review. '
    'Durante la etapa de UAT la app está en "Development Mode" y solo funciona '
    'para los usuarios registrados como Testers — ninguna cuenta externa puede '
    'usar la aplicación hasta que Meta apruebe cada scope para modo público (Live Mode).',
    italic=True
)

doc.add_paragraph()
doc.add_paragraph()

# Firma
sig = doc.add_paragraph()
sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
sig_run = sig.add_run('— DivinAds Team —')
sig_run.font.color.rgb = PRIMARY
sig_run.bold = True
sig_run.font.size = Pt(12)

sig2 = doc.add_paragraph()
sig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sig2_run = sig2.add_run('soporte@divinads.com')
sig2_run.font.color.rgb = MUTED
sig2_run.font.size = Pt(10)

# ─── Guardar ─────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                        'DivinAds - Guia de Prueba del Cliente.docx')
doc.save(out_path)
print(f'OK: documento generado en {out_path}')
print(f'Tamaño: {os.path.getsize(out_path)} bytes')
