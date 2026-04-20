"""
Genera el documento ejecutivo de migración V2 en formato Word (.docx).
Destinatario: cliente / stakeholder no técnico.
Fuente: docs/migration/MIGRATION_V2.md (resumen ejecutivo).
Uso:  python tools/generate_migration_exec.py
Salida: "DivinAds - Plan de Migracion V2 (Ejecutivo).docx"
"""
from __future__ import annotations
import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Paleta ----------
PRIMARY  = RGBColor(0x6B, 0x21, 0xA8)   # morado DivinAds
ACCENT   = RGBColor(0x10, 0xB9, 0x81)
WARN     = RGBColor(0xF5, 0x9E, 0x0B)
DANGER   = RGBColor(0xDC, 0x26, 0x26)
MUTED    = RGBColor(0x6B, 0x72, 0x80)
BG_SOFT  = "F3F4F6"

OUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "DivinAds - Plan de Migracion V2 (Ejecutivo).docx",
)

doc = Document()

# Estilos base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for sect in doc.sections:
    sect.top_margin    = Cm(2.2)
    sect.bottom_margin = Cm(2.2)
    sect.left_margin   = Cm(2.4)
    sect.right_margin  = Cm(2.4)


def shade(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_h(text: str, level: int = 1, color: RGBColor = PRIMARY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 0:
        run.font.size = Pt(26)
    elif level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_p(text: str, bold: bool = False, color: RGBColor | None = None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p


def add_num(text: str):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(2)
    return p


def add_callout(title: str, body: str, color: RGBColor = ACCENT):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    c = tbl.rows[0].cells[0]
    shade(c, "F0FDF4")
    c.paragraphs[0].text = ""
    p1 = c.paragraphs[0]
    r1 = p1.add_run(title + "\n")
    r1.bold = True
    r1.font.color.rgb = color
    r2 = p1.add_run(body)
    r2.font.size = Pt(10)
    doc.add_paragraph()


# =============================== PORTADA ==============================
add_h("DivinAds", level=0, color=PRIMARY)
add_h("Plan de Migración a Arquitectura SaaS Segura (V2)", level=1, color=MUTED)
add_p("Documento ejecutivo para el cliente", color=MUTED)
add_p(f"Versión 2.0  ·  {date.today().strftime('%d de %B de %Y')}", color=MUTED)
doc.add_paragraph()

meta = doc.add_table(rows=4, cols=2)
meta.autofit = True
rows = [
    ("Destinatario",   "Equipo DivinAds / Stakeholders"),
    ("Tipo",           "Plan de migración y hoja de ruta"),
    ("Alcance",        "Arquitectura, seguridad, datos, plazos y entregables"),
    ("Confidencialidad","Interno — no distribuir sin autorización"),
]
for i, (k, v) in enumerate(rows):
    c0 = meta.rows[i].cells[0]
    c1 = meta.rows[i].cells[1]
    c0.text = k; shade(c0, "EDE9FE")
    c1.text = v
    for p in c0.paragraphs:
        for r in p.runs: r.bold = True; r.font.color.rgb = PRIMARY

doc.add_page_break()

# =============================== ÍNDICE ==============================
add_h("Contenido", level=1)
for i, t in enumerate([
    "Resumen ejecutivo",
    "Por qué migrar: situación actual vs. objetivo",
    "Arquitectura objetivo (visión general)",
    "Principios de seguridad",
    "Modelo de datos y aislamiento multi-tenant",
    "Licenciamiento y planes comerciales",
    "Plan de migración por fases",
    "Cronograma y esfuerzo estimado",
    "Riesgos y mitigaciones",
    "Entregables finales",
    "Criterios de aceptación (DoD)",
    "Próximos pasos",
], start=1):
    add_p(f"{i}. {t}")
doc.add_page_break()

# =========================== 1. RESUMEN EJECUTIVO =====================
add_h("1. Resumen ejecutivo", level=1)
add_p(
    "DivinAds migrará de una aplicación monolítica (extensión Chrome + servidor "
    "Node.js local con estado en disco) a una arquitectura SaaS moderna basada en "
    "tres piezas desacopladas: (1) panel web Next.js desplegado en Vercel, "
    "(2) base de datos multi-tenant en Supabase Postgres con Row-Level Security "
    "y Auth integrado, y (3) extensión Chrome MV3 'ligera' que solo ejecuta "
    "tareas autorizadas por el backend. Esta transformación elimina el riesgo de "
    "guardar tokens de Facebook en archivos locales, habilita facturación por "
    "planes, y permite escalar a múltiples agencias y usuarios sin reescribir "
    "el producto."
)
add_callout(
    "Impacto de negocio",
    "• Producto comercializable por suscripción (Stripe-ready).\n"
    "• Aislamiento estricto entre clientes (compliance y confianza).\n"
    "• Cero datos sensibles en el navegador del usuario final.\n"
    "• Auditoría completa de todas las acciones sobre cuentas Meta.",
)

# =========================== 2. POR QUÉ MIGRAR =======================
add_h("2. Por qué migrar: situación actual vs. objetivo", level=1)
add_p("Situación actual (V1):", bold=True)
for b in [
    "Servidor Node.js ejecutándose en el equipo del usuario (localhost:8080).",
    "Tokens OAuth almacenados en tokens.json cifrados con AES-256-GCM, pero persistidos en disco.",
    "Sin multi-tenant: cada instalación es autónoma y no comparte datos.",
    "Sin base de datos central, sin auditoría, sin planes/licencias.",
    "La extensión Chrome contiene toda la lógica del negocio (ofuscada).",
]:
    add_bullet(b)

add_p("Arquitectura objetivo (V2):", bold=True)
for b in [
    "Panel web en Vercel con autenticación Supabase (email + Google).",
    "Base de datos Postgres centralizada con aislamiento por tenant y RLS.",
    "OAuth de Meta 100% server-side; tokens cifrados con pgcrypto.",
    "Extensión Chrome minimalista: solo UI + llamadas al backend.",
    "Emparejamiento seguro extensión ↔ cuenta vía código de 6 dígitos (TTL 5 min).",
    "Licenciamiento por planes (Trial, Starter, Pro, Enterprise) y feature flags.",
    "Auditoría completa de acciones (audit_logs).",
]:
    add_bullet(b)

# =========================== 3. ARQUITECTURA ==========================
add_h("3. Arquitectura objetivo (visión general)", level=1)
add_p(
    "Tres capas claramente separadas. Ninguna información sensible "
    "(tokens, secretos) reside nunca en el cliente."
)
arch = doc.add_table(rows=4, cols=2)
arch.autofit = True
for i, (k, v) in enumerate([
    ("Capa 1 · Cliente",
     "Extensión Chrome MV3 (UI + llamadas HTTPS) + panel web Next.js (dashboards y administración)."),
    ("Capa 2 · Backend",
     "API Routes de Next.js desplegadas como funciones serverless en Vercel. Responsables de OAuth con Meta, pareo de dispositivos, validación de licencias, proxy al Graph API."),
    ("Capa 3 · Datos",
     "Supabase: Postgres con RLS, Auth (usuarios del panel), Storage (avatares/exports) y Vault (claves de cifrado)."),
    ("Integraciones",
     "Meta Graph API v20.0 · Upstash Redis (rate-limit) · Stripe (facturación en fase 2) · Sentry (observabilidad)."),
]):
    c0 = arch.rows[i].cells[0]; c1 = arch.rows[i].cells[1]
    c0.text = k; c1.text = v
    shade(c0, "EDE9FE")
    for p in c0.paragraphs:
        for r in p.runs: r.bold = True; r.font.color.rgb = PRIMARY

# =========================== 4. SEGURIDAD ============================
add_h("4. Principios de seguridad", level=1)
for t, d in [
    ("Zero-trust con el navegador",
     "La extensión jamás ve el token Meta. Todas las llamadas al Graph API pasan por el backend."),
    ("Aislamiento multi-tenant",
     "Row-Level Security en Postgres garantiza que un cliente nunca puede leer datos de otro, ni por error de código."),
    ("Tokens cifrados en reposo",
     "Los access_token de Meta se cifran con pgcrypto pgp_sym_encrypt; la clave vive en Supabase Vault y solo se accede por RPC SECURITY DEFINER."),
    ("OAuth endurecido",
     "State firmado HMAC-SHA256, single-use, TTL 10 min. appsecret_proof en cada llamada al Graph API."),
    ("Pareo de dispositivos",
     "Código 6 dígitos, hash sha256 en DB, TTL 5 min, consumo atómico, rate-limit 20/IP/10 min."),
    ("Auditoría",
     "Cada acción sensible (conexión Meta, revocación, acceso a token) queda registrada en audit_logs."),
    ("Cumplimiento",
     "Cabeceras HTTPS estrictas (HSTS, CSP), rate-limit global, body-limit 100 KB, registros mínimos por 90 días."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(t + ". ")
    r1.bold = True; r1.font.color.rgb = PRIMARY
    p.add_run(d)
    p.paragraph_format.space_after = Pt(2)

# =========================== 5. MODELO DE DATOS =======================
add_h("5. Modelo de datos y aislamiento multi-tenant", level=1)
add_p(
    "Supabase Postgres contiene 11 tablas (ver MIGRATION_V2.md §10). "
    "Las más relevantes para el negocio son:"
)
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 4"
hdr = tbl.rows[0].cells
hdr[0].text = "Tabla"; hdr[1].text = "Propósito"
for t, d in [
    ("tenants",             "Cada agencia/cliente es un tenant (espacio aislado)."),
    ("tenant_members",      "Usuarios y rol (owner/admin/operator/viewer)."),
    ("licenses",            "Plan activo, estado y periodo de facturación."),
    ("meta_connections",    "Cuentas de Meta conectadas al tenant (sin token)."),
    ("meta_tokens",         "Access token cifrado; acceso solo vía RPC."),
    ("oauth_transactions",  "Transacciones OAuth en curso (state firmado)."),
    ("device_pairings",     "Códigos de pareo extensión (hash + TTL)."),
    ("extension_installs",  "Sesiones activas de la extensión Chrome."),
    ("audit_logs",          "Trazabilidad completa."),
]:
    row = tbl.add_row().cells
    row[0].text = t; row[1].text = d

# =========================== 6. LICENCIAMIENTO =======================
add_h("6. Licenciamiento y planes comerciales", level=1)
lic = doc.add_table(rows=1, cols=3)
lic.style = "Light Grid Accent 4"
h = lic.rows[0].cells
h[0].text = "Plan"; h[1].text = "Incluye"; h[2].text = "Límites"
for p, inc, lim in [
    ("Trial",      "Acceso 14 días, módulos BM y Cuentas.",      "1 usuario, 1 cuenta Meta."),
    ("Starter",    "BM + Cuentas + Páginas + Pixels.",           "3 usuarios, 3 cuentas Meta."),
    ("Pro",        "Starter + Advantage+ + Attribution + exports.", "10 usuarios, 10 cuentas Meta."),
    ("Enterprise", "Todo + SSO + SLA + soporte prioritario.",    "Ilimitado (negociado)."),
]:
    r = lic.add_row().cells
    r[0].text = p; r[1].text = inc; r[2].text = lim

# =========================== 7. PLAN POR FASES =======================
add_h("7. Plan de migración por fases", level=1)
fases = [
    ("Fase 0 · Preparación",  "Scaffolding monorepo, cuentas Vercel/Supabase, secretos.",    "2 días"),
    ("Fase 1 · Base de datos","Migraciones SQL, RLS, seeds, claves en Vault.",               "3 días"),
    ("Fase 2 · Auth panel",   "Login email + Google, layouts protegidos, invitaciones.",     "3 días"),
    ("Fase 3 · OAuth Meta",   "Start + callback + refresh + revocación.",                    "4 días"),
    ("Fase 4 · Pareo ext.",   "Endpoints create/redeem, UI en panel y extensión.",           "3 días"),
    ("Fase 5 · Proxy Graph",  "Endpoints BM / Ads / Pages / Pixel detrás del backend.",      "6 días"),
    ("Fase 6 · Extensión V2", "Migrar UI Chrome a llamadas backend, retirar libs.",          "6 días"),
    ("Fase 7 · Licencias",    "Middleware requireActiveLicense + feature flags + Stripe.",   "4 días"),
    ("Fase 8 · QA y release", "Suites E2E, pen-test ligero, onboarding cliente, go-live.",   "4 días"),
]
tf = doc.add_table(rows=1, cols=3)
tf.style = "Light Grid Accent 4"
hh = tf.rows[0].cells
hh[0].text = "Fase"; hh[1].text = "Alcance"; hh[2].text = "Esfuerzo"
for f, a, e in fases:
    r = tf.add_row().cells
    r[0].text = f; r[1].text = a; r[2].text = e

add_callout(
    "Esfuerzo total estimado",
    "Entre 34 y 48 días hábiles de ingeniería, según disponibilidad del equipo.",
    color=PRIMARY,
)

# =========================== 8. CRONOGRAMA ============================
add_h("8. Cronograma sugerido", level=1)
add_p(
    "Asumiendo un equipo de 1 senior full-stack a tiempo completo y un "
    "apoyo part-time de QA, el cronograma semanal es:"
)
for s, w in [
    ("Semana 1", "Fase 0 + Fase 1 (infra y base de datos)."),
    ("Semana 2", "Fase 2 + inicio Fase 3 (auth + OAuth Meta)."),
    ("Semana 3", "Fin Fase 3 + Fase 4 (pareo extensión)."),
    ("Semana 4-5", "Fase 5 (proxy Graph) + inicio Fase 6."),
    ("Semana 6-7", "Fin Fase 6 (extensión V2) + Fase 7 (licencias)."),
    ("Semana 8", "Fase 8 (QA, auditoría final, release a producción)."),
]:
    add_bullet(f"{s}: {w}")

# =========================== 9. RIESGOS ===============================
add_h("9. Riesgos y mitigaciones", level=1)
rk = doc.add_table(rows=1, cols=3)
rk.style = "Light Grid Accent 4"
hh = rk.rows[0].cells
hh[0].text = "Riesgo"; hh[1].text = "Probabilidad / Impacto"; hh[2].text = "Mitigación"
for r, p, m in [
    ("Rechazo de App Review por Meta", "Media / Alto",
     "Pedir el mínimo de scopes, screencast del flujo, usar Business Verification."),
    ("Pérdida de sesión al migrar usuarios V1", "Baja / Medio",
     "Script de importación + comunicación previa con instrucciones de reconexión."),
    ("Bug en políticas RLS expone datos entre tenants", "Baja / Crítico",
     "Pruebas automáticas de RLS por tabla, code review, pentest antes de release."),
    ("Costos inesperados de Vercel/Supabase", "Media / Bajo",
     "Planes de arranque gratuitos + alertas de billing + caché de Graph en Redis."),
    ("Revocación de token del usuario", "Media / Medio",
     "Detección del error, envío de email, re-autenticación guiada en el panel."),
]:
    row = rk.add_row().cells
    row[0].text = r; row[1].text = p; row[2].text = m

# =========================== 10. ENTREGABLES =========================
add_h("10. Entregables finales", level=1)
for e in [
    "Repositorio monorepo (apps/web, apps/extension, packages/*, supabase/, docs/).",
    "Migraciones SQL (0001_init.sql, 0002_rls_policies.sql) listas para Supabase.",
    "Contratos TypeScript + Zod en packages/types.",
    "Endpoints Vercel (/api/meta/*, /api/extension/pair/*, /api/graph/*).",
    "Extensión Chrome V2 minimalista (UI + Bearer token).",
    "Suite de pruebas unitarias + integración + E2E.",
    "Documento maestro MIGRATION_V2.md (23 secciones) y este ejecutivo.",
    "Guía de pruebas para el cliente (UAT).",
    "Runbook de operación (backups, rotación de claves, incidentes).",
]:
    add_bullet(e)

# =========================== 11. DoD =================================
add_h("11. Criterios de aceptación (Definition of Done)", level=1)
for c in [
    "0 vulnerabilidades high/critical en dependencias (npm audit).",
    "Cobertura de tests ≥ 80 % en packages/core y apps/web/api.",
    "100 % de endpoints con validación Zod y manejo de errores normalizado.",
    "RLS habilitado y verificado con tests automáticos en todas las tablas con tenant_id.",
    "OAuth Meta pasa App Review de Facebook (estado 'In development' → 'Live').",
    "Despliegue reproducible: 1 comando (pnpm deploy) para staging y prod.",
    "Documentación operativa revisada y firmada por el cliente.",
]:
    add_bullet(c)

# =========================== 12. PRÓXIMOS PASOS ======================
add_h("12. Próximos pasos inmediatos", level=1)
for i, s in enumerate([
    "Aprobación formal de este plan por parte del cliente.",
    "Creación de proyecto Supabase y cuenta Vercel.",
    "Configuración de secretos en Vercel (FB_APP_ID, FB_APP_SECRET, OAUTH_STATE_SECRET, JWT_SECRET, SUPABASE_*).",
    "Ejecución de migraciones 0001 y 0002 en Supabase (dev + staging).",
    "Kickoff técnico — Fase 0/1.",
], start=1):
    add_num(s)

doc.add_page_break()
add_h("Anexo A — Contacto y seguimiento", level=1)
add_p(
    "Este documento complementa MIGRATION_V2.md (documento técnico maestro de "
    "23 secciones). Cualquier cambio de alcance se gestionará vía pull request "
    "en el repositorio y quedará auditado."
)

doc.save(OUT_PATH)
print(f"Documento generado: {OUT_PATH}")
